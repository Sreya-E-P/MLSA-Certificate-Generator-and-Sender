import os
import pandas as pd
from docx import Document
from docx2pdf import convert

# Function to validate file paths
def validate_file_path(file_path, file_type):
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"The {file_type} file does not exist at the specified path: {file_path}")
    return file_path

def validate_folder_path(folder_path):
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)
    return folder_path

# Prompt user for input
csv_file = input("Enter the path to the CSV file: ").strip()
downloads_folder = input("Enter the path to the downloads folder: ").strip()

# Validate paths
csv_file = validate_file_path(csv_file, "CSV")
downloads_folder = validate_folder_path(downloads_folder)

# Load the CSV file with the mapping of old names (FirstName_LastName) to new names (Name)
df = pd.read_csv(csv_file)

# Load the Word document
docx_file = input("Enter the path to the Word template file: ").strip()
docx_file = validate_file_path(docx_file, "Word template")

# Prompt user for input
Event_or_challenge_name = input("Enter the name of the Event/challenge: ").strip()
mlsa_name = input("Enter the name of MLSA: ").strip()
rank = input("Enter the Milestone of MLSA: ").strip()
date = input("Enter the Date: ").strip()

# Iterate through each student's name in the CSV file
for index, row in df.iterrows():
    student_name = row['Name'].strip()

    # Make a copy of the certificate template for each student
    doc = Document(docx_file)

    # Replace the placeholder text with user-provided values
    placeholders = {
        'FirstName_LastName': student_name,
        'ChallengeName': Event_or_challenge_name,
        'Date': date,
        'MLSAName': mlsa_name,
        'Rank': rank
    }

    for placeholder, value in placeholders.items():
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)

    # Save the edited certificate as a temporary .docx file in the Downloads folder
    temp_docx_filename = os.path.join(downloads_folder, f'{student_name}_Certificate.docx')
    doc.save(temp_docx_filename)

    # Convert the temporary .docx file to PDF using docx2pdf
    pdf_filename = os.path.join(downloads_folder, f'{student_name}_Certificate.pdf')
    convert(temp_docx_filename, pdf_filename)

print("Certificates have been generated in the respective Folder.")

